from docx import Document

# Create a new Word document
doc = Document()
doc.add_heading("📘 Full Documentation: RAG Chatbot with Streamlit + Groq (LLaMA3) + FAISS", 0)

# Section 1: Libraries & Imports
doc.add_heading("1. Libraries & Imports", level=1)
doc.add_paragraph("""
These libraries are used to build the chatbot:
- streamlit: UI framework for building the chatbot app
- os: File/directory management
- dotenv: Load API keys from .env
- time: Measure response time
- langchain_groq: Groq LLM wrapper
- langchain_openai: Optional OpenAI embeddings
- langchain_community.embeddings: HuggingFace embeddings
- langchain_community.vectorstores: FAISS vector database
- langchain_community.document_loaders: Load PDFs
- langchain.text_splitter: Break docs into chunks
- langchain_core.prompts: Prompt templating
- langchain.chains.combine_documents: Stuffing docs into LLM
- langchain.chains: RAG pipeline
""")

doc.add_paragraph("Code:")
doc.add_paragraph("""
import streamlit as st
import os
from dotenv import load_dotenv
import time

from langchain_groq import ChatGroq
from langchain_openai import OpenAIEmbeddings
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_community.document_loaders import PyPDFDirectoryLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_core.prompts import ChatPromptTemplate
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain.chains import create_retrieval_chain
""")

# Section 2: Environment Variables
doc.add_heading("2. Environment Variables", level=1)
doc.add_paragraph("Load the Groq API key from `.env` file:")
doc.add_paragraph("""
load_dotenv()
groq_api_key = os.getenv("Groq_key")  # ✅ Make sure .env has: Groq_key=your_api_key
""")

# Section 3: Initialize LLM
doc.add_heading("3. Initialize LLM (Groq + LLaMA3)", level=1)
doc.add_paragraph("We initialize the Groq-hosted LLaMA3 model:")
doc.add_paragraph("""
llm = ChatGroq(groq_api_key=groq_api_key, model_name="llama-3.1-8b-instant")
""")

# Section 4: Prompt Template
doc.add_heading("4. Prompt Template", level=1)
doc.add_paragraph("This ensures the model answers only from context:")
doc.add_paragraph("""
prompt = ChatPromptTemplate.from_template(\"""
Answer the question based on the provided context only.
Please provide the most accurate response based on the question.

<context>
{context}
</context>

Question: {input}
\""")
""")

# Section 5: Embedding & Vector DB
doc.add_heading("5. Create Embeddings & Vector Database", level=1)
doc.add_paragraph("""
This function loads PDFs, splits them into chunks, embeds them using HuggingFace, 
and stores them in a FAISS vector database.
""")
doc.add_paragraph("""
def create_vectors_embeddings():
    if "vectors" not in st.session_state:
        try:
            # Load PDFs
            data_path = os.path.join(os.getcwd(), "data")
            loader = PyPDFDirectoryLoader(data_path)
            raw_docs = loader.load()
            if not raw_docs:
                st.warning("⚠️ No documents found in the 'data/' folder.")
                return

            # Split into chunks
            splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
            split_docs = splitter.split_documents(raw_docs[:10])  # limit to 10 for speed

            # Embed with HuggingFace
            embeddings = HuggingFaceEmbeddings()
            vectordb = FAISS.from_documents(split_docs, embeddings)

            # Save state
            st.session_state.embeddings = embeddings
            st.session_state.vectors = vectordb
            st.session_state.docs = split_docs

            st.success("✅ Vector database created from documents.")
        except Exception as e:
            st.error(f"❌ Error creating embeddings: {e}")
""")

# Section 6: Streamlit UI Setup
doc.add_heading("6. Streamlit UI Setup", level=1)
doc.add_paragraph("""
We configure the Streamlit app layout with title and description.
""")
doc.add_paragraph("""
st.set_page_config(page_title="📄 RAG Chatbot with Groq + LLaMA3")
st.title("📄 RAG Document Q&A with Groq + LLaMA3")
st.write("Ask a question based on the research papers in the `data/` folder.")
""")

# Section 7: Button to Trigger Embedding Creation
doc.add_heading("7. Button to Trigger Embedding Creation", level=1)
doc.add_paragraph("""
A button is provided to create the vector DB when clicked.
""")
doc.add_paragraph("""
if st.button("📁 Document Embeddings"):
    create_vectors_embeddings()
""")

# Section 8: User Query Input
doc.add_heading("8. User Query Input", level=1)
doc.add_paragraph("Users enter their query here:")
doc.add_paragraph("""
user_prompt = st.text_input("🔍 Enter your query:")
""")

# Section 9: RAG Flow
doc.add_heading("9. RAG Flow: Retrieval + LLM Answering", level=1)
doc.add_paragraph("""
If the user provides a query, the retrieval chain is run. 
The system fetches relevant chunks and passes them to the LLM with the prompt.
""")
doc.add_paragraph("""
if user_prompt:
    if "vectors" not in st.session_state:
        st.warning("⚠️ Please create the vector database first by clicking '📁 Document Embeddings'.")
    else:
        try:
            document_chain = create_stuff_documents_chain(llm, prompt)
            retriever = st.session_state.vectors.as_retriever()
            retrieval_chain = create_retrieval_chain(retriever, document_chain)

            start_time = time.process_time()
            response = retrieval_chain.invoke({"input": user_prompt})
            duration = time.process_time() - start_time

            st.markdown(f"⏱️ **Response time:** {duration:.2f} seconds")
            st.write("🤖 **Answer:**")
            st.success(response.get("answer", "No answer returned."))

            with st.expander("📚 Document Chunks Used for Answer"):
                for i, doc in enumerate(response.get("context", [])):
                    st.markdown(f"**Chunk {i+1}:**")
                    st.write(doc.page_content)

        except Exception as e:
            st.error(f"❌ Error during retrieval: {e}")
""")

# Workflow Summary
doc.add_heading("✅ Workflow Summary", level=1)
doc.add_paragraph("""
1. Upload PDFs → stored in data/
2. Click “📁 Document Embeddings” → create FAISS DB
3. Enter query → RAG pipeline retrieves docs + sends to LLaMA3
4. Model answers only from context
5. User sees answer + supporting chunks
""")

# Save document
output_path = "/mnt/data/RAG_Chatbot_With_Code_Documentation.docx"
doc.save(output_path)

output_path
